import uuid
from typing import List, Dict, Any
from datetime import datetime
import PyPDF2
from docx import Document
import tiktoken
import csv
import openpyxl
import io


class DocumentService:
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = upload_dir
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

    async def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process uploaded document and extract text"""
        text = await self._extract_text(file_path, filename)
        chunks = self._chunk_text(text)

        document_id = str(uuid.uuid4())

        return {
            "document_id": document_id,
            "filename": filename,
            "chunks": chunks,
            "chunk_count": len(chunks),
            "upload_date": datetime.now().isoformat(),
        }

    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from different file formats"""
        if filename.endswith(".pdf"):
            return self._extract_pdf_text(file_path)
        elif filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif filename.endswith(".csv"):
            return self._extract_csv_text(file_path)
        elif filename.endswith(".xlsx"):
            return self._extract_xlsx_text(file_path)
        elif filename.endswith(".docx"):
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV file"""
        text = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                text.append(", ".join(row))
        return "\n".join(text)

    def _extract_xlsx_text(self, file_path: str) -> str:
        """Extract text from XLSX file"""
        text = []
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for sheet in workbook.worksheets:
            text.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                # Filter out None values and convert to string
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    text.append(", ".join(row_text))
        return "\n".join(text)

    async def extract_text_from_bytes(self, file_data: bytes, filename: str) -> str:
        """Extract text from file bytes (for in-memory processing)"""
        if filename.endswith(".pdf"):
            # PyPDF2 can read from BytesIO
            pdf_file = io.BytesIO(file_data)
            text = ""
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text

        elif filename.endswith(".txt"):
            return file_data.decode("utf-8")

        elif filename.endswith(".csv"):
            text_data = file_data.decode("utf-8")
            f = io.StringIO(text_data)
            reader = csv.reader(f)
            lines = [", ".join(row) for row in reader]
            return "\n".join(lines)

        elif filename.endswith(".xlsx"):
            xlsx_file = io.BytesIO(file_data)
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True, data_only=True)
            text = []
            for sheet in workbook.worksheets:
                text.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) for cell in row if cell is not None]
                    if row_text:
                        text.append(", ".join(row_text))
            return "\n".join(text)

        elif filename.endswith(".docx"):
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])

        else:
            raise ValueError(f"Unsupported file type: {filename}")

    def _chunk_text(
        self, text: str, chunk_size: int = 512, overlap: int = 50
    ) -> List[Dict[str, Any]]:
        """Split text into chunks with overlap"""
        tokens = self.tokenizer.encode(text)
        chunks = []

        for i in range(0, len(tokens), chunk_size - overlap):
            chunk_tokens = tokens[i : i + chunk_size]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(
                {"text": chunk_text, "start_index": i, "token_count": len(chunk_tokens)}
            )

        return chunks
