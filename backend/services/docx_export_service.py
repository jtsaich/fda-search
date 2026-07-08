import re
from io import BytesIO
from typing import Optional

from docx import Document

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
DEFAULT_FILENAME = "fda-search-answer.docx"
_MAX_FILENAME_STEM_LENGTH = 80
_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.+)$")
_UNSAFE_FILENAME_RE = re.compile(r"[^a-z0-9._-]+")


class DocxSource:
    def __init__(
        self,
        title: Optional[str] = None,
        filename: Optional[str] = None,
        score: Optional[float] = None,
        text: Optional[str] = None,
    ) -> None:
        self.title = title
        self.filename = filename
        self.score = score
        self.text = text


def sanitize_docx_filename(title: Optional[str]) -> str:
    if not title or not title.strip():
        return DEFAULT_FILENAME

    stem = title.strip().lower().encode("ascii", "ignore").decode("ascii")
    stem = _UNSAFE_FILENAME_RE.sub("_", stem).strip("._-")
    stem = re.sub(r"_+", "_", stem)[:_MAX_FILENAME_STEM_LENGTH].strip("._-")
    if not stem:
        return DEFAULT_FILENAME
    return f"{stem}.docx"


def build_docx_bytes(
    *,
    content: str,
    title: Optional[str] = None,
    sources: Optional[list[DocxSource]] = None,
) -> bytes:
    document = Document()

    if title and title.strip():
        document.add_heading(title.strip(), level=0)

    _add_markdownish_content(document, content)

    referenced_sources = [source for source in sources or [] if _source_has_content(source)]
    if referenced_sources:
        document.add_heading("Sources Referenced", level=1)
        for source in referenced_sources:
            _add_source(document, source)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_markdownish_content(document: Document, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            document.add_heading(heading_match.group(2).strip(), level=level)
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1).strip(), style="List Bullet")
            continue

        numbered_match = _NUMBERED_RE.match(line)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1).strip(), style="List Number")
            continue

        document.add_paragraph(line)


def _add_source(document: Document, source: DocxSource) -> None:
    label = source.title or source.filename or "Untitled source"
    document.add_paragraph(label, style="List Bullet")

    if source.filename and source.filename != label:
        document.add_paragraph(f"File: {source.filename}")
    if source.score is not None:
        document.add_paragraph(f"Score: {source.score:g}")
    if source.text and source.text.strip():
        document.add_paragraph(source.text.strip())


def _source_has_content(source: DocxSource) -> bool:
    return any(
        value is not None and (not isinstance(value, str) or bool(value.strip()))
        for value in (source.title, source.filename, source.score, source.text)
    )
